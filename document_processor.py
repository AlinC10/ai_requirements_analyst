import os.path

import docx
# If you want to use the RapidOCR for images just remove the comments from the import and image_parser in
# PyMuPDF4LLMLoader.
# from langchain_community.document_loaders.parsers import RapidOCRBlobParser
from langchain_core.documents import Document
from langchain_pymupdf4llm import PyMuPDF4LLMLoader
from markitdown import MarkItDown
from pymupdf4llm.ocr import rapidocr_api

from doc_splitter import DocSplitter

default_important_metadata_categories: list[str] = ['title', 'author', 'subject', 'keywords', 'source', "chat_id"]
page_break_delimiter = "\n---PAGE BREAK---\n"


def check_file_extension(file_path: str) -> str:
    """Check if the file extension is accepted by the RAG system and return the extension if it is accepted, else throw
    an error which will be caught in the app.py to show an error on the UI.
    Accepted documents: DOCX, PDF, TXT.

    :param file_path: Document file path that contains the extension
    :type file_path: str

    :return: The extension extracted from the file path, if the extension is in the accepted extensions list,
    else it will throw an error.
    :rtype: str
    """

    _, extension = os.path.splitext(file_path)

    # remove . from extension
    # e.g.: .txt -> txt
    extension = extension[1:]

    accepted_types = ["docx", "pdf", "txt"]

    if extension in accepted_types:
        return extension

    raise TypeError(f"Accepted document types: {", ".join(accepted_types)}")


def extract_file_name(file_path: str) -> str:
    """
    Extract the fie name from a given file path.
    e.g.: pdf/file_name.pdf = file_name

    :param file_path: Document's location on disk.
    :type file_path: str
    :return: Only the name of the file, without the path and extension.
    :rtype: str
    """

    base_name = os.path.basename(file_path)
    name_without_ext = os.path.splitext(base_name)[0]

    return name_without_ext


def extract_file_title(file_path: str, text: str = "", meta_title: str | None = None) -> str:
    """
    Finds the best possible title using a fallback chain.

    :param file_path: Document's location on disk.
    :type file_path: str
    :param text: Text content of the document.
    :type text: str
    :param meta_title: Title obtained from the document metadata.
    :type meta_title: str
    :return: The title, which could be in the following order: meta_title, the header 1 from the text Markdown or the
    file name.
    :rtype: str
    """

    if meta_title and meta_title.strip() not in ["", "Untitled"]:
        return meta_title

    # search for "# " which is used for titles in Markdown format
    for line in text.splitlines():
        line = line.strip()
        if line.startswith('# '):
            return line[2:].strip(" *")

    # create the title from file path, removing "_" and "-" symbols
    # e.g.: pdf/file_name.pdf = File Name
    file_name_without_ext = extract_file_name(file_path)
    clean_name = file_name_without_ext.replace("_", " ").replace("-", " ").title()

    return clean_name


def remove_unused_pdf_metadata(docs: Document,
                               important_metadata_categories: list[str] | None = None) -> Document:
    """
    Remove unused information from metadata of the PDF files.

    :param docs: The PDF document which will have the unused metadata removed.
    :type docs: Document
    :param important_metadata_categories: The metadata that should not be removed.
    :type important_metadata_categories: list[str] | None
    :return: The PDF document with the removed metadata.
    :rtype: Document
    """

    if important_metadata_categories is None:
        important_metadata_categories = default_important_metadata_categories

    keys_to_delete = [key for key in docs.metadata if key not in important_metadata_categories]

    for key in keys_to_delete:
        del docs.metadata[key]

    return docs


def split_doc(document: Document, extension: str, chunk_size: int = 1000, chunk_overlap: int = 100,
              separators: list[str] | None = None) -> list[Document]:
    """
    Split the documents based on the rules defined in the DocSplitter class.

    :param document: Document that will be split in multiple chunks depending on his extension.
    :type document: Document
    :param extension: Used for determine the best approach to split the document.
    :type extension: str
    :param chunk_size: Maximum number of characters for a text chunk. Default = 1000 characters.
    :type chunk_size: int
    :param chunk_overlap: Maximum number of characters for the overlap between 2 chunks. Default = 100 characters.
    :type chunk_overlap: int
    :param separators: List of characters or Regex expression that will be used for separating the chunks.
    Default = [\"\\n\\n\", \"\\n\", \"(?<=[.?!])\", \" \", \"\"]
    :type separators: list[str] | None
    :return: Document divided into multiple chunks using the rules from above.
    :rtype: list[Document]
    """

    splitter = DocSplitter(extension, chunk_size, chunk_overlap, separators)

    split_document = splitter.document_split(document)

    return split_document


def add_pdf_page_numbers(docs: list[Document]) -> list[Document]:
    """
    Determine the number of each page and add it to the metadata. (only available for PDF files)

    :param docs: Split document that will have his page break delimiter count to determine the page number for each
    chunk.
    :type docs: list[Document]
    :return: The same split document with the page number added to the metadata.
    :rtype: list[Document]
    """

    page_number = 1

    for doc in docs:
        if page_break_delimiter in doc.page_content:
            page_number += 1

            doc.metadata['page'] = page_number

    return docs


def load_doc(file_path: str, chat_id: str, file_name: str) -> list[Document]:
    """
    Determine extension used and load document using the functions made for each extension and divide the document
    into chunks that will be loaded into Chroma Database.

    :param file_path: Used for loading the document. Cannot be used as file name because this is the temporary
    file path, created when it was added in the UI by the user and will not help the citations if it's used as
    file name.
    :type file_path: str
    :param chat_id: Used for better retrieval when the chat conversation history will be added.
    :type chat_id: str
    :param file_name: Original file name, used as source name.
    :type file_name: str

    :return: Document divided into multiple chunks.
    :rtype: list[Document]
    """

    # Checks file extension
    extension = check_file_extension(file_path)

    if extension == 'docx':
        docs = _load_docx(file_path, file_name, chat_id)
    elif extension == "pdf":
        docs = _load_pdf(file_path, file_name, chat_id)
    else:
        docs = _load_txt(file_path, file_name, chat_id)

    split_docs = split_doc(docs, extension)

    if extension == "pdf":
        split_docs = add_pdf_page_numbers(split_docs)

    return split_docs


def _load_docx(file_path: str, file_name: str, chat_id: str) -> Document:
    """
    Load the .docx file using MarkItDown library.

    :param file_path: Used for loading the document. Cannot be used as file name because this is the temporary
    file path, created when it was added in the UI by the user and will not help the citations if it's used as
    file name.
    :type file_path: str
    :param file_name: Original file name, used as source name.
    :type file_name: str
    :param chat_id: Used for better retrieval when the chat conversation history will be added.
    :type chat_id: str

    :return: DOCX document parsed as a Markdown.
    :rtype: Document
    """

    md = MarkItDown()
    document = md.convert(file_path).text_content

    # load .docx properties
    doc = docx.Document(file_path)
    props = doc.core_properties

    # extract the text of the first paragraph styled as a 'Title' from .docx to retrieve the document title
    title = None
    try:
        for para in doc.paragraphs[:5]:
            if para.style.name.startswith("Title") and para.text.strip():
                title = para.text.strip()
                break
    except Exception:
        pass

    metadata = {"author": props.author or "Unknown",
                "title": title or extract_file_title(file_name, document, props.title),
                "category": props.category or "Unknown",
                "source": file_name,
                "chat_id": chat_id
                }

    return Document(
        page_content=document,
        metadata=metadata
    )


def _load_pdf(file_path: str, file_name: str, chat_id: str) -> Document:
    """
    Load the PDF file into the RAM using PyMuPDF4LLMLoader with RapidOCR for retrieving text from images.

    :param file_path: Used for loading the document. Cannot be used as file name because this is the temporary
    file path, created when it was added in the UI by the user and will not help the citations if it's used as
    file name.
    :type file_path: str
    :param file_name: Original file name, used as source name.
    :type file_name: str
    :param chat_id: Used for better retrieval when the chat conversation history will be added.
    :type chat_id: str

    :return: PDF document parsed as a Markdown.
    :rtype: Document
    """

    important_metadata_categories = default_important_metadata_categories

    loader = PyMuPDF4LLMLoader(
        file_path=file_path,
        mode="single",
        pages_delimiter=page_break_delimiter,
        extract_images=False,

        # use RapidOCR for extracted image fragments
        # images_parser=RapidOCRBlobParser(),

        # tell PyMuPDF4LLM to use RapidOCR instead of Tesseract for scanned pages
        ocr_function=rapidocr_api.exec_ocr
    )

    document = loader.load()[0]
    document.metadata.update({"chat_id": chat_id, "source": file_name})

    title = extract_file_title(file_name, document.page_content, document.metadata.get('title', None))
    document.metadata.update({'title': title})

    document = remove_unused_pdf_metadata(document, important_metadata_categories)

    return document


def _load_txt(file_path: str, file_name: str, chat_id: str) -> Document:
    """Load the .txt file.
    If Python encounters a corrupted or weird character that isn't valid UTF-8, instead of crashing the entire
    script with a UnicodeDecodeError, it simply replaces the bad character with a standard placeholder symbol (
    usually ``).

    :param file_path: Used for loading the document. Cannot be used as file name because this is the temporary
    file path, created when it was added in the UI by the user and will not help the citations if it's used as
    file name.
    :type file_path: str
    :param file_name: Original file name, used as source name.
    :type file_name: str
    :param chat_id: Used for better retrieval when the chat conversation history will be added.
    :type chat_id: str

    :return: TXT document loaded as string and converted to a Document with metadata added.
    :rtype: Document
    """
    with open(file_path, mode="r", encoding="utf-8", errors="replace") as file:
        txt = file.read()
        return Document(page_content=txt,
                        metadata={"source": file_name, "title": extract_file_title(file_name), "chat_id": chat_id})


def format_for_llm(docs: list[Document]) -> str:
    """
    Format the retrieved chunks that will be sent to the LLM as a string, where every chunk has it's source,
    title (and page if it contains that) metadata into the chunks for citations.

    :param docs: Retrieved chunks that will be sent to the LLM
    :type docs: list[Document]

    :return: Chunks joined togheter with the metadata added to every chunk for citations.
    :rtype: str
    """

    formated_context = "\n\n".join(
        f"""[Source: {doc.metadata.get("source", "Unknown")} | Title: {doc.metadata.get('title', 'Unknown')} 
{f"| Page: {doc.metadata['page']}" if doc.metadata.get("page", False) else ""}]\n
{doc.page_content}"""
        for doc in docs
    )

    return formated_context
